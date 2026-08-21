"""Genera el informe académico de la Práctica 2 en formato DOCX.

El archivo se construye desde datos reproducibles guardados en ``data``. De esa
forma el documento puede regenerarse sin editar manualmente resultados ni
perder la relación con el código que los produjo.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS_PATH = ROOT / "data" / "resultados_referencia.json"
OUTPUT_DIR = ROOT / "docs" / "informe"
OUTPUT_PATH = OUTPUT_DIR / "Informe_Practica_2_Daniel_Barillas_22193.docx"
REPOSITORY_URL = (
    "https://github.com/DanielBarillasM/"
    "Practica-2_Daniel-Barillas_22193_Modelacion-y-Simulacion_Sec-30"
)

GREEN = RGBColor(8, 120, 63)
DARK_GREEN = RGBColor(6, 77, 45)
GRAY = RGBColor(90, 103, 95)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Agrega un hipervínculo real mediante los elementos XML de Word.

    ``python-docx`` no expone una operación de alto nivel para enlaces externos.
    Por eso se crea la relación, se asocia su identificador al elemento
    ``w:hyperlink`` y se aplica color y subrayado al texto visible.
    """

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "08783F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, fill: str) -> None:
    """Aplica a una celda el color de fondo hexadecimal indicado.

    La modificación se realiza sobre ``w:tcPr`` porque el sombreado tampoco está
    disponible como propiedad pública completa en ``python-docx``.
    """

    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_number(paragraph) -> None:
    """Inserta en el pie el campo dinámico PAGE reconocido por Word.

    No se escribe un número fijo: Word evalúa los marcadores de inicio,
    instrucción y fin al abrir o actualizar el documento generado.
    """

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Página ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def add_equation(document: Document, text: str) -> None:
    """Añade una expresión matemática centrada con tipografía apropiada.

    Las expresiones del informe son texto matemático legible, no imágenes. Esto
    conserva la posibilidad de seleccionar, editar y buscar su contenido.
    """

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)


def add_result_table(document: Document, results: dict) -> None:
    """Construye la tabla resumen a partir del JSON reproducible.

    Cada fila contrasta una observación simulada con su referencia analítica o
    esperanza. El encabezado y las filas alternas reciben estilos directos para
    mantener una lectura clara aun cuando cambie el tema de Word.
    """

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Ejercicio", "Configuración", "Resultado", "Referencia"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        set_cell_shading(cell, "08783F")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    rows = [
        ("1", "1,000 variables", f"Media {results['ejercicio_1']['media_estimada']:.8f}", f"{results['ejercicio_1']['media_exacta']:.8f}"),
        ("4", "50,000 meses", f"P = {results['ejercicio_4']['probabilidad_estimada']:.4%}", f"{results['ejercicio_4']['probabilidad_referencia']:.4%}"),
        ("5", "10,000 normales", f"μ={results['ejercicio_5']['media']:.5f}; s²={results['ejercicio_5']['varianza']:.5f}", "μ=0; σ²=1"),
        ("6", "λ=2, T=10", f"{results['ejercicio_6']['eventos_observados']} eventos", f"E[N]= {results['ejercicio_6']['eventos_esperados']:.0f}"),
        ("7", "T=10", f"{results['ejercicio_7']['global_eventos']} global; {results['ejercicio_7']['mejorado_eventos']} mejorado", f"E[N]={results['ejercicio_7']['eventos_esperados']:.4f}"),
        ("8", "λ=1, R=5", f"{results['ejercicio_8']['puntos_observados']} puntos", f"E[N]={results['ejercicio_8']['puntos_esperados']:.4f}"),
        ("9", "10,000 normales", f"μ={results['ejercicio_9']['media']:.5f}; s²={results['ejercicio_9']['varianza']:.5f}", "μ=0; σ²=1"),
        ("10", "λ=1, R=2", f"{results['ejercicio_10']['puntos_observados']} puntos", f"E[N]={results['ejercicio_10']['puntos_esperados']:.4f}"),
    ]
    for row_index, row_values in enumerate(rows, 1):
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if row_index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "EAF5EE")


def build_document() -> Document:
    """Ensambla el informe académico completo y devuelve el documento DOCX.

    El proceso carga una única fuente de resultados, define márgenes y estilos,
    crea portada, teoría, desarrollo, tabla de resultados, auditoría y cierre.
    Devolver el objeto separa la construcción de su persistencia en disco y
    facilita reutilizar o probar el contenido desde otro módulo.
    """

    # La corrida de referencia se lee antes de crear párrafos para que todas las
    # cifras incluidas en el informe provengan del mismo archivo reproducible.
    results = json.loads(RESULTS_PATH.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Title", 25, DARK_GREEN),
        ("Heading 1", 17, DARK_GREEN),
        ("Heading 2", 13, GREEN),
        ("Heading 3", 11, GREEN),
    ]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color

    header = section.header.paragraphs[0]
    header.text = "Universidad del Valle de Guatemala · CC2017"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.color.rgb = GRAY
    header.runs[0].font.size = Pt(8)
    add_page_number(section.footer.paragraphs[0])

    # Portada
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("UNIVERSIDAD DEL VALLE DE GUATEMALA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GREEN
    for line in [
        "Facultad de Ingeniería",
        "Ciencia de la Computación y Tecnologías de la Información",
        "CC2017 — Modelación y Simulación",
        "Ciclo 2, 2026 · Sección 30",
    ]:
        p = document.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = GRAY

    document.add_paragraph()
    title = document.add_paragraph("Práctica 2")
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Generación de Variables Aleatorias Continuas")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(17)
    subtitle.runs[0].font.color.rgb = GREEN

    document.add_paragraph()
    for label, value in [
        ("Elaborado por", "Pablo Daniel Barillas Moreno"),
        ("Carné", "22193"),
        ("Sección", "30"),
        ("Semilla reproducible", "22193"),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    document.add_paragraph()
    repo = document.add_paragraph()
    repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    repo.add_run("Repositorio que presenta este informe:\n").bold = True
    add_hyperlink(repo, REPOSITORY_URL, REPOSITORY_URL)

    document.add_page_break()

    document.add_heading("1. Introducción", level=1)
    document.add_paragraph(
        "Esta práctica desarrolla métodos para generar variables aleatorias continuas y procesos "
        "de Poisson. El trabajo combina la respuesta teórica de cada inciso con implementaciones "
        "auditables en Python y una interfaz Streamlit que conserva los números aleatorios, presenta "
        "tablas completas y permite comparar resultados simulados con referencias matemáticas."
    )

    document.add_heading("2. Objetivos", level=1)
    for text in [
        "Aplicar transformación inversa, composición y aceptación-rechazo.",
        "Generar normales estándar mediante rechazo exponencial y el método polar.",
        "Simular procesos de Poisson homogéneos, no homogéneos y bidimensionales.",
        "Cuantificar la incertidumbre Monte Carlo y validar los algoritmos con teoría.",
        "Presentar una aplicación reproducible, documentada y organizada por responsabilidades.",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("3. Metodología y reproducibilidad", level=1)
    document.add_paragraph(
        "Se utilizó NumPy con el generador PCG64. Elegí 22193, mi número de carné, como semilla "
        "predeterminada. La semilla inicializa el generador y permite repetir las tablas; no significa "
        "que se emplee el mismo número en todos los lanzamientos. La interfaz permite cambiarla."
    )

    exercises = [
        (
            "4.1 Ejercicio 1 — Exponencial condicionada",
            "La CDF condicional se invierte directamente para evitar el alto desperdicio que tendría un método de rechazo.",
            "X = -ln(1 - U(1-e^(-0.05)))",
            "La media exacta es 1 - 0.05/(e^0.05-1) = 0.02479167535.",
        ),
        (
            "4.2 Ejercicio 2 — Método de composición",
            "Se genera I con probabilidades p_i y, condicionado a I=i, se genera X con CDF F_i. La probabilidad total produce la mezcla solicitada.",
            "P(X≤x) = Σ p_i F_i(x)",
            "La aplicación permite editar pesos y observar la CDF teórica y empírica.",
        ),
        (
            "4.3 Ejercicio 3 — Aplicaciones de composición",
            "Los incisos se expresan como mezclas de CDF potencia; el inciso (b) combina una exponencial de tasa 2 con una uniforme(0,1).",
            "Para F_i(x)=x^i se usa X=U^(1/i)",
            "Las tres soluciones incluyen selección de componente, tabla y gráfica de CDF.",
        ),
        (
            "4.4 Ejercicio 4 — Cartera de seguros",
            "El número de reclamaciones es binomial y la suma condicionada es Gamma. Se estima la probabilidad de superar $50,000.",
            "P(S>c) = Σ P(N=n) P(Gamma(n,800)>c)",
            "La referencia matemática es 0.1070977013.",
        ),
        (
            "4.5 Ejercicio 5 — Normal por rechazo exponencial",
            "Se implementa el Ejemplo 5f con dos exponenciales, signo equiprobable y reciclaje del residual exponencial independiente.",
            "Aceptar si Y2 > (Y1-1)^2/2",
            "La aceptación teórica es sqrt(π/(2e)) ≈ 0.76017; la eficiencia esperada es 1.64 exponenciales y 1.32 cuadrados por normal.",
        ),
        (
            "4.6 Ejercicio 6 — Poisson homogéneo",
            "Los tiempos entre eventos son exponenciales de tasa λ y se acumulan hasta superar T.",
            "N(T) ~ Poisson(λT)",
            "La UI presenta la trayectoria escalonada N(t) y el uniforme de cada llegada.",
        ),
        (
            "4.7 Ejercicio 7 — Poisson no homogéneo",
            "Se usa adelgazamiento con M=7. La mejora divide el horizonte en intervalos y usa cotas locales decrecientes.",
            "E[N(10)] = 30 + 4 ln(11) = 39.59158109",
            "La corrida de referencia reduce las propuestas de 72 a 41.",
        ),
        (
            "4.8 Ejercicio 8 — Poisson bidimensional",
            "El conteo es Poisson con media λπR² y los puntos se distribuyen uniformemente en el círculo.",
            "r=R√U1, θ=2πU2",
            "Para λ=1 y R=5 se esperan 25π ≈ 78.5398 puntos.",
        ),
        (
            "4.9 Ejercicio 9 — Método polar",
            "Marsaglia evita seno y coseno: acepta pares dentro del círculo unitario y produce dos normales independientes.",
            "X=V1√(-2 ln(S)/S), Y=V2√(-2 ln(S)/S)",
            "La aceptación teórica es π/4 ≈ 0.7854.",
        ),
        (
            "4.10 Ejercicio 10 — Poisson bidimensional: teoría",
            "Para toda región A, N(A) es Poisson con media λ|A| y regiones disjuntas tienen conteos independientes.",
            "N(A) ~ Poisson(λ|A|)",
            "Se incluye un ejemplo numérico completo para λ=1 y R=2.",
        ),
    ]
    statements = [
        "Sea X exponencial con media 1. Genere eficientemente 1,000 valores condicionados a X<0.05, estime su media y determine el valor exacto.",
        "Explique cómo generar una variable cuya CDF es F(x)=Σ p_i F_i(x), con pesos no negativos que suman 1.",
        "Usando composición, proporcione algoritmos para las distribuciones de los incisos (a), (b) y (c) indicadas en el PDF.",
        "Con 1,000 asegurados, probabilidad de reclamación 0.05 y montos exponenciales de media $800, estime P(S>$50,000).",
        "Genere variables normales con el rechazo exponencial de tasa 1 presentado en el Ejemplo 5f.",
        "Genere las primeras T unidades de tiempo de un proceso de Poisson homogéneo con tasa λ.",
        "Genere por adelgazamiento el proceso con λ(t)=3+4/(t+1) en [0,10] y proponga una mejora.",
        "Genere y grafique un proceso de Poisson bidimensional en un círculo para λ=1 y R=5.",
        "Explique el método polar, su utilidad frente a Box–Muller y desarrolle un ejemplo numérico completo.",
        "Defina el proceso de Poisson bidimensional, explique sus aplicaciones y desarrolle un ejemplo para λ=1 y R=2.",
    ]
    document.add_heading("4. Desarrollo de los ejercicios", level=1)
    for statement, (title_text, explanation, equation, result_text) in zip(statements, exercises):
        document.add_heading(title_text, level=2)
        statement_paragraph = document.add_paragraph()
        statement_paragraph.add_run("Enunciado: ").bold = True
        statement_paragraph.add_run(statement)
        document.add_paragraph(explanation)
        add_equation(document, equation)
        document.add_paragraph(result_text)

    document.add_heading("5. Resultados de la corrida reproducible", level=1)
    document.add_paragraph(
        "Los siguientes valores se obtuvieron con la semilla 22193. Las diferencias respecto de una "
        "esperanza o referencia son variabilidad normal de Monte Carlo y no implican un error."
    )
    add_result_table(document, results)

    document.add_heading("6. Aplicación y organización del repositorio", level=1)
    document.add_paragraph(
        "La interfaz está en app/, los algoritmos en src/practica2/, las pruebas en tests/, los "
        "resultados reproducibles en data/, los documentos en docs/ y la identidad visual en assets/. "
        "Esta separación evita mezclar presentación con cálculo y facilita mantener el proyecto."
    )
    p = document.add_paragraph()
    p.add_run("Repositorio: ").bold = True
    add_hyperlink(p, REPOSITORY_URL, REPOSITORY_URL)

    document.add_heading("7. Ejecución", level=1)
    document.add_paragraph("Instalar dependencias:")
    add_equation(document, "python -m pip install -r requirements.txt")
    document.add_paragraph("Iniciar la aplicación:")
    add_equation(document, "streamlit run app/app.py")
    document.add_paragraph("Ejecutar pruebas:")
    add_equation(document, "python -m pytest -q")

    document.add_heading("8. Auditoría contra las instrucciones", level=1)
    document.add_paragraph(
        "La revisión se realizó contra los diez ejercicios y todos sus incisos del PDF de la "
        "Práctica 2 de CC2017. No se encontró una rúbrica separada aplicable a esta actividad. "
        "La rúbrica localizada en Descargas corresponde a Construcción de Compiladores y, por "
        "tratarse de otro curso, no se utilizó."
    )
    document.add_paragraph(
        "Como resultado de la auditoría se incorporó el enunciado en cada pantalla, se optimizó "
        "el Ejercicio 5 reciclando la exponencial residual independiente y se cubrió el extremo "
        "U=0 en la inversión del conteo Poisson espacial."
    )

    document.add_heading("9. Conclusiones", level=1)
    for conclusion in [
        "La transformación inversa evita el rechazo ineficiente en la exponencial truncada.",
        "El método de composición convierte mezclas de CDF en algoritmos simples y verificables.",
        "Los generadores normales reproducen adecuadamente media, varianza y tasas de aceptación teóricas.",
        "Las cotas locales mejoran de manera clara el adelgazamiento del proceso no homogéneo.",
        "La transformación radial R√U es indispensable para uniformidad espacial en un círculo.",
        "La separación entre teoría, cálculo e interfaz produce un repositorio reproducible y mantenible.",
    ]:
        document.add_paragraph(conclusion, style="List Bullet")

    document.add_section(WD_SECTION.NEW_PAGE)
    return document


def main() -> None:
    """Genera el DOCX final en la ruta oficial del repositorio.

    La carpeta se crea si aún no existe. Guardar se mantiene fuera de
    ``build_document`` para que construir el objeto no produzca efectos laterales
    cuando el módulo sea importado por herramientas o pruebas.
    """

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"Informe generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
